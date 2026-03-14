import io
from docling.document_converter import DocumentConverter
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.backend.pypdf_backend import PyPdfDocumentBackend


class DocumentProcessor:
    def __init__(self):
        self._converter = None
    
    @property
    def converter(self) -> DocumentConverter:
        if self._converter is None:
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True
            pipeline_options.do_table_structure = True
            
            self._converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(
                        pipeline_options=pipeline_options,
                        backend=PyPdfDocumentBackend,
                    )
                }
            )
        return self._converter
    
    def process(self, file_content: bytes, filename: str) -> str:
        """Process document and return text content."""
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        try:
            if file_ext == 'pdf':
                return self._process_pdf(file_content)
            elif file_ext in ['doc', 'docx']:
                return self._process_doc(file_content, file_ext)
            elif file_ext == 'txt':
                return self._process_txt(file_content)
            else:
                return f"[Unsupported format: {file_ext}]"
        except Exception as e:
            return f"[Error processing document: {str(e)}]"
    
    def _process_pdf(self, file_content: bytes) -> str:
        doc = self.converter.convert(io.BytesIO(file_content))
        return doc.document.export_to_markdown()
    
    def _process_doc(self, file_content: bytes, file_ext: str) -> str:
        try:
            import docx
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_ext}') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            try:
                if file_ext == 'docx':
                    doc = docx.Document(tmp_path)
                    text = '\n'.join([para.text for para in doc.paragraphs])
                else:
                    text = self._extract_from_doc(tmp_path)
                return text
            finally:
                os.unlink(tmp_path)
        except Exception as e:
            return f"[Error processing DOC: {str(e)}]"
    
    def _extract_from_doc(self, doc_path: str) -> str:
        try:
            from docx import Document
            doc = Document(doc_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text
        except:
            return "[Unable to extract text from .doc. Please convert to .docx]"
    
    def _process_txt(self, file_content: bytes) -> str:
        return file_content.decode('utf-8', errors='ignore')


document_processor = DocumentProcessor()
